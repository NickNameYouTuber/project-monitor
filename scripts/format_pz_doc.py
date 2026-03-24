import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


VALID_H2 = {
    "Описание предметной области и области применения",
    "Описание целевого рынка и потенциальных пользователей",
    "Описание существующих аналогов",
    "Описание важности разработки программного продукта",
    "Хронология разработки программного продукта",
}

SPECIAL_H1 = {
    "Содержание",
    "Введение",
    "Заключение",
    "Список литературы",
    "Список использованных источников",
}


def iter_docx_files():
    root = Path(__file__).resolve().parents[1]
    for path in root.glob("ПЗ_4П1_9_22_Хорлин_НВ*.docx"):
        yield path


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def remove_numbering(paragraph):
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def remove_style_numbering(style):
    element = style.element
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        return
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_run_font(run, size_pt=14, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def ensure_single_run(paragraph, text, size_pt=14, bold=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold)
    return run


def normalize_title(text):
    text = text.replace("\xa0", " ").strip()
    text = re.sub(r"^[•\-]\s*", "", text)
    text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text)
    return text.strip()


def is_code_paragraph(text):
    stripped = text.strip()
    if not stripped:
        return False
    code_starts = (
        "//",
        "@",
        "public ",
        "private ",
        "protected ",
        "const ",
        "export ",
        "<LiveKitRoom",
        "}",
        "return ",
        "if (",
        "switch (",
        "void ",
    )
    return "\n" in text and stripped.startswith(code_starts)


def set_update_fields(doc):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)
    else:
        existing.set(qn("w:val"), "true")


def add_toc_field(paragraph):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = Cm(0)

    begin = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    begin._r.append(fld_char)
    set_run_font(begin, size_pt=14)

    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'
    instr._r.append(instr_text)
    set_run_font(instr, size_pt=14)

    separate = paragraph.add_run()
    sep_char = OxmlElement("w:fldChar")
    sep_char.set(qn("w:fldCharType"), "separate")
    separate._r.append(sep_char)
    set_run_font(separate, size_pt=14)

    placeholder = paragraph.add_run("Обновите содержание в Word при открытии документа.")
    set_run_font(placeholder, size_pt=14)

    end = paragraph.add_run()
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end._r.append(end_char)
    set_run_font(end, size_pt=14)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    clear_paragraph(paragraph)
    run = paragraph.add_run()
    set_run_font(run, size_pt=12)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def set_page_number_start(section, start=2):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def format_paragraph(paragraph, kind):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)

    if kind == "content-title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.line_spacing = 1.5
        ensure_single_run(paragraph, "Содержание", size_pt=14, bold=True)
        return

    if kind == "h1":
        remove_numbering(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        fmt.line_spacing = 2.0
        fmt.page_break_before = True
        for run in paragraph.runs:
            set_run_font(run, size_pt=14, bold=True)
        return

    if kind == "h2":
        remove_numbering(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        fmt.line_spacing = 2.0
        for run in paragraph.runs:
            set_run_font(run, size_pt=14, bold=True)
        return

    if kind == "table-number":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.line_spacing = 1.5
        for run in paragraph.runs:
            set_run_font(run, size_pt=14)
        return

    if kind == "table-title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.line_spacing = 1.5
        for run in paragraph.runs:
            set_run_font(run, size_pt=14)
        return

    if kind == "figure-caption":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.line_spacing = 1.5
        for run in paragraph.runs:
            set_run_font(run, size_pt=14)
        return

    if kind == "source":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.line_spacing = 1.5
        for run in paragraph.runs:
            set_run_font(run, size_pt=14)
        return

    if kind == "code":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.line_spacing = 1.5
        for run in paragraph.runs:
            set_run_font(run, size_pt=14)
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.left_indent = Cm(0)
    fmt.first_line_indent = Cm(1.5)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run, size_pt=14)


def normalize_document(path: Path):
    doc = Document(str(path))

    remove_style_numbering(doc.styles["Heading 1"])
    remove_style_numbering(doc.styles["Heading 2"])

    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.footer_distance = Cm(1.27)
    set_page_number_start(section, start=2)

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_number(footer_para)
    set_update_fields(doc)

    paragraphs = doc.paragraphs
    intro_index = None
    for index, paragraph in enumerate(paragraphs):
        if normalize_title(paragraph.text) == "Введение":
            intro_index = index
            break

    if intro_index is not None and intro_index > 2:
        for paragraph in reversed(list(doc.paragraphs[2:intro_index])):
            delete_paragraph(paragraph)

    paragraphs = doc.paragraphs
    if paragraphs:
        paragraphs[0].style = doc.styles["Normal"]
        format_paragraph(paragraphs[0], "content-title")

    if len(paragraphs) > 1:
        paragraphs[1].style = doc.styles["Normal"]
        add_toc_field(paragraphs[1])

    chapter = 0
    subchapter = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text
        base = normalize_title(text)

        if base == "Содержание":
            paragraph.style = doc.styles["Normal"]
            format_paragraph(paragraph, "content-title")
            continue

        if base in SPECIAL_H1 or base.startswith("Приложение") or paragraph.style.name == "Heading 1":
            if base == "Введение" or base in SPECIAL_H1 or base.startswith("Приложение"):
                paragraph.style = doc.styles["Heading 1"]
                ensure_single_run(paragraph, base, size_pt=14, bold=True)
                format_paragraph(paragraph, "h1")
                continue

        if base in VALID_H2:
            paragraph.style = doc.styles["Heading 2"]
            if chapter == 0:
                chapter = 1
            subchapter += 1
            ensure_single_run(paragraph, f"{chapter}.{subchapter}. {base}", size_pt=14, bold=True)
            format_paragraph(paragraph, "h2")
            continue

        if base and paragraph.style.name == "Heading 1" and base not in SPECIAL_H1:
            chapter += 1
            subchapter = 0
            paragraph.style = doc.styles["Heading 1"]
            ensure_single_run(paragraph, f"{chapter}. {base}", size_pt=14, bold=True)
            format_paragraph(paragraph, "h1")
            continue

        if paragraph.style.name == "Heading 2":
            paragraph.style = doc.styles["Normal"]

        stripped = text.strip()
        if re.fullmatch(r"Таблица\s+\d+\.", stripped):
            format_paragraph(paragraph, "table-number")
        elif stripped.startswith("Продолжение таблицы"):
            format_paragraph(paragraph, "table-title")
        elif stripped.startswith("Рисунок "):
            format_paragraph(paragraph, "figure-caption")
        elif stripped.startswith("Источник:"):
            format_paragraph(paragraph, "source")
        elif is_code_paragraph(text):
            format_paragraph(paragraph, "code")
        else:
            format_paragraph(paragraph, "normal")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Введение":
            paragraph.paragraph_format.page_break_before = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    for run in paragraph.runs:
                        set_run_font(run, size_pt=12)

    duplicate_marker = "Подраздел составлен на основании бланка задания на ВКР"
    duplicate_indices = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip().startswith(duplicate_marker)
    ]
    if len(duplicate_indices) > 1:
        start = duplicate_indices[1]
        for paragraph in list(doc.paragraphs[start:]):
            delete_paragraph(paragraph)

    doc.save(str(path))


def main():
    for path in iter_docx_files():
        normalize_document(path)
        print(f"Formatted {path.name}")


if __name__ == "__main__":
    main()