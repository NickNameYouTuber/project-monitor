from pathlib import Path

from docx import Document


def dump_doc_structure(doc_path: Path) -> None:
    document = Document(doc_path)
    print(f"Документ: {doc_path.name}")
    print(f"Параграфов: {len(document.paragraphs)}")
    print(f"Таблиц: {len(document.tables)}")
    for table_index, table in enumerate(document.tables):
        print(f"\nТаблица {table_index + 1}:")
        for row_index, row in enumerate(table.rows):
            row_cells = []
            for cell_index, cell in enumerate(row.cells):
                text = cell.text.replace("\n", " | ").strip()
                row_cells.append(f"[{cell_index}] {text}")
            joined_cells = " || ".join(row_cells)
            print(f"  Строка {row_index}: {joined_cells}")


if __name__ == "__main__":
    doc_relative_path = Path("project-monitor") / "Бланк задания на ВКР.docx"
    dump_doc_structure(doc_relative_path)
