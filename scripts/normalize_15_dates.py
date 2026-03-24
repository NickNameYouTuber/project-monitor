from pathlib import Path

from docx import Document


REPLACEMENT = "11.12.2025 - 23.02.2026 — "


def rewrite_paragraph(paragraph) -> bool:
    text = paragraph.text
    if not text.startswith("• "):
        return False
    if len(text) < 16:
        return False
    if not (text[2:4].isdigit() and text[5:7].isdigit() and text[8:12].isdigit()):
        return False
    if text[4] != "." or text[7] != ".":
        return False

    dash_pos = text.find("–")
    if dash_pos == -1:
        dash_pos = text.find("—")
    if dash_pos == -1:
        dash_pos = text.find("-")
    if dash_pos == -1:
        return False

    updated = REPLACEMENT + text[dash_pos + 1 :].strip()

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    for path in root.glob("ПЗ_4П1_9_22_Хорлин_НВ*.docx"):
        if path.name.startswith("~$"):
            continue
        doc = Document(str(path))
        changed = 0
        for paragraph in doc.paragraphs:
            if rewrite_paragraph(paragraph):
                changed += 1
        if changed:
            doc.save(str(path))
        print(f"{path.name}: {changed}")


if __name__ == "__main__":
    main()